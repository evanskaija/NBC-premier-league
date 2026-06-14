import os
import sys
import subprocess

# Ensure python-docx is installed
try:
    import docx
except ImportError:
    print("python-docx not found. Installing...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    """Set cell background color."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def add_heading(doc, text, level):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.bold = True
    
    if level == 1:
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0, 53, 102) # Dark Blue
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 180, 216) # Teal
    else:
        run.font.size = Pt(11.5)
        run.font.color.rgb = RGBColor(50, 50, 50)
    return p

def add_body(doc, text, bold=False, italic=False, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(51, 51, 51)
    return p

def add_bullet(doc, text, bold_prefix=""):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    
    if bold_prefix:
        r_bold = p.add_run(bold_prefix)
        r_bold.font.name = 'Arial'
        r_bold.font.size = Pt(11)
        r_bold.font.bold = True
        r_bold.font.color.rgb = RGBColor(51, 51, 51)
        
    r_text = p.add_run(text)
    r_text.font.name = 'Arial'
    r_text.font.size = Pt(11)
    r_text.font.color.rgb = RGBColor(51, 51, 51)
    return p

def main():
    doc = Document()
    
    # Page Setup - 1 inch margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Document Header Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(24)
    title.paragraph_format.space_after = Pt(6)
    run_title = title.add_run("TANZANIA FOOTBALL HUB WEBSITE")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(26)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0, 53, 102)

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(36)
    run_sub = subtitle.add_run("Project Status & System Specification Report")
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(14)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(120, 120, 120)

    # 1. INTRODUCTION
    add_heading(doc, "1. Introduction", level=1)
    
    add_heading(doc, "Background", level=2)
    add_body(doc, "Football is the most popular sport in Tanzania, but information about leagues, clubs, players, fixtures, results, transfers, and statistics is often scattered across different platforms.")
    
    add_heading(doc, "Problem Statement", level=2)
    add_body(doc, "Football fans, clubs, players, journalists, and sponsors face difficulties accessing accurate and centralized football information in Tanzania.")

    add_heading(doc, "Proposed Solution", level=2)
    add_body(doc, "Develop a comprehensive Tanzania Football Hub website that provides live scores, fixtures, results, standings, player statistics, team profiles, transfers, news, videos, galleries, and football-related information in one platform.")

    add_heading(doc, "Objectives", level=2)
    add_heading(doc, "Main Objective", level=3)
    add_body(doc, "To develop a modern web-based football information management platform for Tanzanian football.")
    add_heading(doc, "Specific Objectives", level=3)
    add_bullet(doc, " Provide live football scores.", bold_prefix="Live Scores:")
    add_bullet(doc, " Display fixtures and results.", bold_prefix="Fixtures & Results:")
    add_bullet(doc, " Manage team and player information.", bold_prefix="Manage Team & Player Information:")
    add_bullet(doc, " Track transfers and statistics.", bold_prefix="Track Transfers & Statistics:")
    add_bullet(doc, " Publish football news and media.", bold_prefix="Publish News & Media:")
    add_bullet(doc, " Improve fan engagement.", bold_prefix="Improve Fan Engagement:")

    # 2. LITERATURE REVIEW
    add_heading(doc, "2. Literature Review", level=1)
    add_body(doc, "To design a top-tier platform, the system was benchmarked against leading global and local sports websites:")
    add_bullet(doc, " Flashscore & Sofascore", bold_prefix="Existing football websites:")
    add_bullet(doc, " Strong in real-time scores and data coverage but lack local contextual content (like official regulations, governance directories, and local media feeds).", bold_prefix="Strengths and weaknesses:")
    add_bullet(doc, " HTML5, Vanilla CSS3, Vanilla JavaScript, local storage, custom Swahili translation engine, and Python automation helper scripts.", bold_prefix="Technologies used in sports websites:")
    add_bullet(doc, " Local sports fans are highly active on mobile devices and crave visual, fast-loading, media-heavy content like video highlights and interactive polls.", bold_prefix="Research findings:")

    # 3. SYSTEM ANALYSIS
    add_heading(doc, "3. System Analysis", level=1)
    
    add_heading(doc, "Existing System", level=2)
    add_body(doc, "Challenges in the current Tanzanian sports information landscape include:")
    add_bullet(doc, "Information is scattered across different social media channels (Instagram, X, YouTube) and personal blogs.")
    add_bullet(doc, "Delayed updates for standings, player statistics, and match records.")
    add_bullet(doc, "Lack of a centralized, digitized database of official tournament rules (Kanuni) and executive board contacts.")
    add_bullet(doc, "Tanzanian sports platforms often lack mobile responsiveness, causing poor navigation on smartphone screens.")

    add_heading(doc, "Proposed System", level=2)
    add_body(doc, "Benefits of the newly developed Tanzania Football Hub website:")
    add_bullet(doc, "One single digital location for all fixtures, results, tables, transfers, news, and rulebooks.")
    add_bullet(doc, "Extremely fast navigation and mobile-first responsive templates built with clean Vanilla JS/CSS.")
    add_bullet(doc, "A gorgeous, immersive visual experience leveraging modern UI design rules (dark mode, glassmorphism, dynamic countdowns, animated reveal grids).")
    add_bullet(doc, "Dual-language support (English and Swahili) that translates the entire interface instantly.")

    # 4. SYSTEM REQUIREMENTS
    add_heading(doc, "4. System Requirements", level=1)
    
    add_heading(doc, "Functional Requirements", level=2)
    add_bullet(doc, "Users can search for players, matches, clubs, or articles via a unified search box.")
    add_bullet(doc, "Users can toggle the site language between English and Kiswahili on-the-fly.")
    add_bullet(doc, "Users can toggle the website between Premium Dark and Light modes.")
    add_bullet(doc, "Matches filterable by league, date, or club with Day, Week, and Month calendars.")
    add_bullet(doc, "Matches detail pages feature real-time scores, team line-ups, match momentum charts, text commentary, live ratings, and fan voting polls.")
    add_bullet(doc, "Standings tables automatically display positions, form guides, and highlight qualification zones (CAF Champions League, Confederation Cup, Relegation).")
    add_bullet(doc, "The hub directory catalogs teams, detailed player profile cards, market values, and transfer window rumor bars.")
    add_bullet(doc, "A digital library displays tournament rules with built-in PDF viewing panes.")
    add_bullet(doc, "Fan zone featuring multiple-choice trivia quizzes.")

    add_heading(doc, "Non-Functional Requirements", level=2)
    add_bullet(doc, "Secure links for member portal authentication.")
    add_bullet(doc, "Persistent header, footer, and navigation menus across all pages.")
    add_bullet(doc, "Optimized CSS styles and Vanilla JS to guarantee fast loading under mobile data connections.")
    add_bullet(doc, "Full responsiveness on mobile, tablet, and desktop screens with a slide-out navigation menu for mobile screens.")
    add_bullet(doc, "Highly intuitive interface with prominent call-to-actions, clear indicators (LIVE, UPCOMING, FT), and consistent typography.")

    # 5. SYSTEM ARCHITECTURE & FILE STRUCTURE
    add_heading(doc, "5. System Architecture & File Structure", level=1)
    add_body(doc, "The system follows a lightweight, serverless client-side architecture styled with modern responsive rules. Global elements are managed via a centralized Python pipeline.")
    
    add_heading(doc, "Key Components", level=2)
    add_bullet(doc, "Outfit typography, global variable sheets for themes, glassmorphism containers, animated scroll-reveals, and podium cards.", bold_prefix="CSS Design System (style.css):")
    add_bullet(doc, "Manages sidebar filtering, theme toggles, mobile drawer triggers, countdown timer mathematics, and the English-to-Swahili translation dictionary.", bold_prefix="JavaScript Logic (script.js):")
    add_bullet(doc, "Iterates through all 31+ HTML files to automatically replace outdated headers/footers and inject standardized templates with active tab states.", bold_prefix="Layout Automation Script (standardize.py):")

    add_heading(doc, "File Directory Catalog", level=2)
    add_bullet(doc, "Portal home dashboard with sliders, reels, and tickers.", bold_prefix="index.html:")
    add_bullet(doc, "TPLB History, Board Members, and Secretariat Tree.", bold_prefix="about.html:")
    add_bullet(doc, "Interactive quiz center.", bold_prefix="quiz.html:")
    add_bullet(doc, "Match lists and live commentaries.", bold_prefix="fixtures.html, results.html, livescores.html, calendar.html:")
    add_bullet(doc, "Squad lists, player profile metrics, and coach biographies.", bold_prefix="teams.html, team-details.html, players.html, player-details.html, coach-details.html:")
    add_bullet(doc, "podiums for top scorers, assists, clean sheets, and form timelines.", bold_prefix="statistics.html:")
    add_bullet(doc, "transfer window status, confirmed deals, and confidence ratings.", bold_prefix="transfers.html:")
    add_bullet(doc, "news, features, videos, match reports, and gallery assets.", bold_prefix="news.html, news-details.html, features.html, reports.html, videos.html, gallery.html:")
    add_bullet(doc, "division-specific pages, including FA Cup and continental tournaments.", bold_prefix="nbcpremierleague.html, championship.html, firstleague.html, womensleague.html, facup.html, cafmatches.html:")
    add_bullet(doc, "official league rules with PDF frames.", bold_prefix="kanuni.html:")

    # 6. CURRENT STATUS & GAP ANALYSIS
    add_heading(doc, "6. Current Status & Gap Analysis", level=1)
    add_body(doc, "The website has established a gorgeous, fully-responsive frontend (approximately 75% complete). The following table outlines the remaining work needed for 100% production launch:")

    # Table
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Light Shading Accent 1'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Module / Component'
    hdr_cells[1].text = 'Current Status (Done)'
    hdr_cells[2].text = 'Remaining Gaps (To Do)'
    
    # Format Header
    for cell in hdr_cells:
        set_cell_background(cell, "003566")
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
                
    row_data = [
        ("Dynamic Data Integration", "HTML tables and player metrics display clean static placeholders.", "Connect pages to a backend database or external sports API to serve dynamic squads, stats, and records."),
        ("Live Scores Engine", "Match center renders momentum bars, commentary grids, and player ratings.", "Implement WebSockets or database polling to stream real-time match events and active scores."),
        ("Quiz System", "Quiz center displays four games with premium graphics and cards.", "Add JavaScript gameplay logic, score counting, time limits, and database leaderboard sheets."),
        ("Member Portal & CMS", "Links direct users to external login systems.", "Construct a local administrator panel to allow writers to post news and upload match reports directly.")
    ]
    
    for i, (comp, status, gap) in enumerate(row_data):
        row_cells = table.rows[i+1].cells
        row_cells[0].text = comp
        row_cells[1].text = status
        row_cells[2].text = gap
        
        # Zebra striping
        if i % 2 == 1:
            for cell in row_cells:
                set_cell_background(cell, "F2F7FA")
                
    doc.add_paragraph("\n")

    # 7. CONCLUSION
    add_heading(doc, "7. Conclusion", level=1)
    add_body(doc, "The Tanzania Football Hub is a state-of-the-art sports portal designed to modernize how football data is consumed in Tanzania. By providing a centralized, mobile-responsive, and dual-language platform, it solves the challenge of scattered information and lack of digitized governance. Moving from static layouts to a dynamic database-driven backend represents the final major phase of this development.")

    # Save to user desktop/project directory
    output_path = r"c:\Users\EVANS\OneDrive\Desktop\my projects\Nbc pl\Tanzania_Football_Hub_Report.docx"
    doc.save(output_path)
    print(f"Word document saved successfully to: {output_path}")

if __name__ == "__main__":
    main()
